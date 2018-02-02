import tkinter.filedialog

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH

from models.reference import Reference
from utils.apa_widgets import *


def save_as_docx():
    try:
        input_file_name = tkinter.filedialog.asksaveasfilename(defaultextension='.docx',
                                                               filetypes=[
                                                                   ("Word Documents", '*.docx')
                                                               ])
        if input_file_name:
            create_word_doc(input_file_name)
    except PermissionError:
        print("File open elsewhere!")


def create_word_doc(filename):
    doc = docx.Document()
    doc.add_paragraph('References').alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in Reference.get_sorted_references():
        p1 = f"{r.contributors} ({r.journal.year_published}). {r.journal.article_title}. "

        if r.journal.volume_issue != '' or r.journal.pages != '':
            j_title = r.journal.journal_title + ", "
        else:
            j_title = r.journal.journal_title + ". "

        if r.journal.volume != '' and r.journal.issue != '':
            j_volume = f"{r.journal.volume}"
        elif r.journal.volume != "" and (r.journal.pages == '' or r.journal.doi == ''):
            j_volume = f"{r.journal.volume}."
        else:
            j_volume = ''

        if r.journal.issue != '' and r.journal.pages != '':
            j_issue = f"({r.journal.issue}), "
        elif r.journal.issue != '' and r.pages == '':
            j_issue = f"({r.journal.issue})."
        else:
            j_issue = ''

        if r.journal.page_start != '' and r.journal.page_end != '':
            j_pages = f"{r.journal.page_start}-{r.journal.page_end}."
        elif r.journal.page_start != '' and r.journal.page_end == '':
            j_pages = f"{r.journal.page_start}."
        else:
            j_pages = ''

        if r.journal.doi != '':
            j_doi = f" {r.journal.doi}"
        else:
            j_doi = ''

        ref = doc.add_paragraph(p1)
        ref.add_run(j_title).italic = True
        ref.add_run(j_volume).italic = True
        ref.add_run(j_issue)
        ref.add_run(j_pages)
        ref.add_run(j_doi)

    doc.save(filename)
