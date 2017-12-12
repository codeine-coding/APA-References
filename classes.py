import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH


class Contributor:
    contributors = []

    def __init__(self, first_name, middle_initial, last_name):
        contributor = "{},{}.{mi}".format(last_name.title(),
                                          first_name[0].title(),
                                          mi=(middle_initial[0] + '.').title() if middle_initial != '' else '')
        self.contributors.append(contributor)

    @classmethod
    def get_contributors(cls):
        contributors = sorted(cls.contributors)

        if len(contributors) == 1:
            authors = contributors[0]
        elif len(contributors) == 2:
            authors = ', & '.join(contributors)
        elif 3 <= len(contributors) <= 7:
            authors = (', '.join(contributors[0:-1]) + ', & {}'.format(contributors[-1]))
        elif len(contributors) > 7:
            authors = (', '.join(contributors[0:6]) + ',...{}'.format(contributors[-1]))
        else:
            authors = ''

        return authors


class Journal:
    def __init__(self, article_title, journal_title, year_published, page_start,
                 page_end, volume, issue, doi_or_website):
        self.article_title = article_title
        self.journal_title = journal_title
        self.year_published = year_published
        self.page_start = page_start
        self.page_end = page_end
        self.volume = volume
        self.issue = issue
        self.doi = doi_or_website

        self.pages = f"{page_start}-{page_end}." if page_start != '' and page_end != '' else ''
        self.volume_issue = ''
        if volume != '' and issue != '':
            self.volume_issue = f"{volume}({issue}),"
        elif volume != "" and issue == '':
            self.volume_issue = f"{volume},"
        else:
            self.volume_issue = ''
        self.doi_or_website = doi_or_website if doi_or_website != '' else ''

        # Last, F. M., Jr. (year). Article Title. Journal Title, Volume(Issue), series, P-Start-P-End.
        # doi:DOI_OR_WEBSITE

    def get_journal_entry(self):
        journal = "({}). {}. {},{} {} {}".format(self.year_published, self.article_title, self.journal_title,
                                                 self.volume_issue, self.pages, self.doi_or_website)
        if journal[-1] == ',':
            journal = journal[0:-1] + '.'
        return journal


class Reference:
    references = []
    formatted_references = []

    def __init__(self, contributors, journal):
        self.contributors = contributors
        self.journal = journal
        self.references.append(self)
        self.formatted_reference = f"{self.contributors} {self.journal.get_journal_entry()}"
        self.formatted_references.append(self.formatted_reference)

    @classmethod
    def get_sorted_references(cls):
        return sorted(cls.references, key=lambda x: x.contributors)

    @classmethod
    def create_word_doc(cls, filename):
        doc = docx.Document()
        for r in cls.get_sorted_references():
            p1 = f"{r.contributors} ({r.journal.year_published}). {r.journal.article_title}. "

            if r.journal.volume_issue != '' or r.journal.pages != '':
                j_title = r.journal.journal_title + ", "
            else:
                j_title = r.journal.journal_title + ". "

            if r.journal.volume != '' and r.journal.issue != '':
                j_volume = f"{r.journal.volume}"
            elif r.journal.volume != "" and (r.journal.pages == '' or r.journal.doi == ''):
                j_volume = f"{r.journal.volume}."
            else:
                j_volume = ''

            if r.journal.issue != '' and r.journal.pages != '':
                j_issue = f"({r.journal.issue}), "
            elif r.journal.issue != '' and r.pages == '':
                j_issue = f"({r.journal.issue})."
            else:
                j_issue = ''

            if r.journal.page_start != '' and r.journal.page_end != '':
                j_pages = f"{r.journal.page_start}-{r.journal.page_end}."
            elif r.journal.page_start != '' and r.journal.page_end == '':
                j_pages = f"{r.journal.page_start}."
            else:
                j_pages = ''

            if r.journal.doi != '':
                j_doi = f" {r.journal.doi}"
            else:
                j_doi = ''
            doc.add_paragraph('Work Cited').alignment = WD_ALIGN_PARAGRAPH.CENTER
            ref = doc.add_paragraph(p1)
            ref.add_run(j_title).italic = True
            ref.add_run(j_volume).italic = True
            ref.add_run(j_issue)
            ref.add_run(j_pages)
            ref.add_run(j_doi)

        doc.save(filename)
